import io
from datetime import datetime
from typing import List

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class PDFExport(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 16)
        self.cell(0, 10, 'Todo App - Task Report', border=False, ln=True, align='C')
        self.set_font('Helvetica', 'I', 10)
        self.cell(0, 5, f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}', ln=True, align='C')
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')

    def chapter_title(self, title):
        self.set_font('Helvetica', 'B', 12)
        self.set_fill_color(99, 102, 241)
        self.set_text_color(255, 255, 255)
        self.cell(0, 8, f' {title}', ln=True, fill=True)
        self.set_text_color(0, 0, 0)
        self.ln(3)

    def todo_item(self, todo):
        status = 'Done' if todo.get('completed') else 'Pending'
        priority = todo.get('priority', 'medium').upper()
        due = todo.get('due_date', 'No due date')
        if due and isinstance(due, str):
            try:
                dt = datetime.fromisoformat(due.replace('Z', '+00:00'))
                due = dt.strftime('%Y-%m-%d')
            except:
                pass

        self.set_font('Helvetica', 'B', 11)
        self.cell(0, 6, f"{todo.get('title', 'Untitled')} [{status}]", ln=True)
        self.set_font('Helvetica', '', 9)
        self.cell(60, 5, f'Priority: {priority}', ln=False)
        self.cell(0, 5, f'Due: {due}', ln=True)

        desc = todo.get('description', '')
        if desc:
            self.set_font('Helvetica', '', 9)
            self.multi_cell(0, 4, f'  {desc[:200]}')
        self.ln(2)

        subtasks = todo.get('subtasks', [])
        if subtasks:
            self.set_font('Helvetica', 'I', 9)
            self.cell(0, 5, '  Subtasks:', ln=True)
            self.set_font('Helvetica', '', 9)
            for sub in subtasks:
                sub_status = 'Done' if sub.get('completed') else 'Open'
                self.cell(0, 4, f"    - {sub.get('title', 'Untitled')} [{sub_status}]", ln=True)
            self.ln(2)


def generate_pdf(todos: List[dict], username: str = "User") -> bytes:
    if not HAS_FPDF:
        raise RuntimeError("fpdf2 is not installed. Run: pip install fpdf2")

    pdf = PDFExport()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 6, f'User: {username}', ln=True)
    pdf.cell(0, 6, f'Total Tasks: {len(todos)}', ln=True)
    completed = sum(1 for t in todos if t.get('completed'))
    pdf.cell(0, 6, f'Completed: {completed}', ln=True)
    pdf.cell(0, 6, f'Pending: {len(todos) - completed}', ln=True)
    pdf.ln(5)

    pending = [t for t in todos if not t.get('completed')]
    done = [t for t in todos if t.get('completed')]

    if pending:
        pdf.chapter_title(f'Pending Tasks ({len(pending)})')
        for todo in pending:
            pdf.todo_item(todo)

    if done:
        pdf.add_page()
        pdf.chapter_title(f'Completed Tasks ({len(done)})')
        for todo in done:
            pdf.todo_item(todo)

    return bytes(pdf.output(dest='S'))


def generate_docx(todos: List[dict], username: str = "User") -> bytes:
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    title = doc.add_heading('Todo App - Task Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}')
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()

    doc.add_paragraph(f'User: {username}').bold = True
    completed = sum(1 for t in todos if t.get('completed'))
    doc.add_paragraph(f'Total Tasks: {len(todos)}  |  Completed: {completed}  |  Pending: {len(todos) - completed}')
    doc.add_paragraph()

    pending = [t for t in todos if not t.get('completed')]
    done = [t for t in todos if t.get('completed')]

    if pending:
        doc.add_heading(f'Pending Tasks ({len(pending)})', level=1)
        for todo in pending:
            _add_todo_to_docx(doc, todo)

    if done:
        doc.add_page_break()
        doc.add_heading(f'Completed Tasks ({len(done)})', level=1)
        for todo in done:
            _add_todo_to_docx(doc, todo)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _add_todo_to_docx(doc, todo: dict):
    title_text = todo.get('title', 'Untitled')
    status = 'Done' if todo.get('completed') else 'Pending'
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}  [{status}]')
    run.bold = True
    run.font.size = Pt(11)

    priority = todo.get('priority', 'medium').upper()
    due = todo.get('due_date', 'No due date')
    if due and isinstance(due, str):
        try:
            dt = datetime.fromisoformat(due.replace('Z', '+00:00'))
            due = dt.strftime('%Y-%m-%d')
        except:
            pass

    p2 = doc.add_paragraph()
    p2.add_run(f'Priority: ').bold = True
    p2.add_run(priority)
    p2.add_run('    |    ')
    p2.add_run('Due: ').bold = True
    p2.add_run(str(due))

    desc = todo.get('description', '')
    if desc:
        p3 = doc.add_paragraph(desc[:300])
        p3.style.font.size = Pt(9)

    subtasks = todo.get('subtasks', [])
    if subtasks:
        p4 = doc.add_paragraph()
        p4.add_run('Subtasks:').bold = True
        for sub in subtasks:
            sub_status = 'Done' if sub.get('completed') else 'Open'
            doc.add_paragraph(f'    - {sub.get("title", "Untitled")} [{sub_status}]', style='List Bullet')

    doc.add_paragraph()
